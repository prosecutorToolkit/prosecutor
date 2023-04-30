import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = docx.Document()

# Configuración de la sección
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(14)
section.page_height = Inches(8.5)

# Configuración del título
titulo = doc.add_paragraph()
titulo_run = titulo.add_run('PROSECUTOR REPORT')
titulo_run.bold = True
titulo.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

# Agregar tabla1
tabla1 = doc.add_table(rows=1, cols=2)
tabla1.style = 'Table Grid'
tabla1.alignment = WD_TABLE_ALIGNMENT.CENTER
tabla1.autofit = False
ancho_columna1 = Inches(3.0)
ancho_columna2 = Inches(14.0 - section.left_margin.inches - section.right_margin.inches - ancho_columna1.inches)
tabla1.columns[0].width = ancho_columna1
tabla1.columns[1].width = ancho_columna2
tabla1.rows[0].height = Inches(2.0)
tabla1.rows[0].cells[0].vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
tabla1.rows[0].cells[1].vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
tabla1.rows[0].cells[0].paragraphs[0].add_run().add_picture('logo.jpg', width=Inches(2.5))
tabla1.rows[0].cells[1].add_paragraph('Descripción del documento').alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

color = 'cfe2f3' # RGB hex format

for row in tabla1.rows:
    for cell in row.cells:
        tcPr = cell._element.tcPr
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

# Para cada celda de la tabla, crear un elemento tcPr y establecer el color de fondo
for row in tabla1.rows:
    for cell in row.cells:
        tcPr = cell._element.tcPr
        tcBack = OxmlElement('w:shd')
        tcBack.set(qn('w:fill'), color)
        tcPr.append(tcBack)

doc.add_paragraph()

# Configuración de la tabla2
tabla2 = doc.add_table(rows=0, cols=7)
tabla2.style = 'Table Grid'
tabla2.autofit = False
ancho_predeterminado = 14.0 - section.left_margin.inches - section.right_margin.inches
ancho_columna0 = Inches(0.4) # ID
ancho_columna1 = Inches(1.0) # NAME
ancho_columna2 = Inches(1.5) # PATH
ancho_columna3 = Inches(1.0) # HASH SHA-256
ancho_columna4 = Inches(0.85) # MATCH
ancho_columna5 = Inches(5.0) # TEXT
ancho_columna6 = Inches(ancho_predeterminado - ancho_columna0.inches - ancho_columna1.inches - ancho_columna2.inches - ancho_columna3.inches - ancho_columna4.inches - ancho_columna5.inches) # METADATA
tabla2.columns[0].width = ancho_columna0
tabla2.columns[1].width = ancho_columna1
tabla2.columns[2].width = ancho_columna2
tabla2.columns[3].width = ancho_columna3
tabla2.columns[4].width = ancho_columna4
tabla2.columns[5].width = ancho_columna5
tabla2.columns[6].width = ancho_columna6

# Agregar los valores a la tabla2
tupla_valores = (('1', 'Juan', '/desktop', '342343fjjws54', 'colombia', 'en colombia hay merluza', 'Foto sacada en Medellin'),
    ('1', 'Juan', '/desktop', '342343fjjws54', 'colombia', 'en colombia hay merluza', 'Foto sacada en Medellin')
    )


headTable2 = ("ID", "NAME", "PATH", "HASH SHA-256", "MATCH", "TEXT", "METADATA")
row = tabla2.add_row().cells
for i in range(7):
    p = row[i].add_paragraph()
    p.add_run(headTable2[i]).bold = True


for row in tupla_valores:
    cells = tabla2.add_row().cells
    for i in range(7):
        cells[i].text = row[i]

for section in doc.sections:
    footer = section.footer
    pagenum = footer.paragraphs[0].add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    pagenum._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    pagenum._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    pagenum._r.append(fldChar)

# for paragraph in titulo:
#     paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER # centrado


for row in tabla1.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = docx.shared.RGBColor(0, 0, 0) # blanco transparente

doc.save('nombre_del_archivo.docx')
